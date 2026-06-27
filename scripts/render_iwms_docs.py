"""IWMS 测试文档渲染脚本。

从 markdown 源文件渲染为多种目标格式：
    PDF (fpdf)、DOCX (python-docx)、XLSX (openpyxl)、HTML、MD、TXT

源文件结构：scripts/test_docs/iwms/source/ 下的 .md 文件，以 YAML-like frontmatter 声明目标格式。

用法:
    cd ai-service
    uv run python ../scripts/render_iwms_docs.py
"""

import os
import re
from pathlib import Path

# ---- 格式渲染器 ----

def _find_chinese_font() -> str | None:
    candidates = [
        # Windows (WSL) — TrueType fonts most reliable with fpdf2
        "/mnt/c/Windows/Fonts/simhei.ttf",
        "/mnt/c/Windows/Fonts/simsunb.ttf",
        "/mnt/c/Windows/Fonts/simfang.ttf",
        "/mnt/c/Windows/Fonts/simkai.ttf",
        # Linux
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/usr/share/fonts/truetype/arphic/ukai.ttc",
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        # project
        str(Path(__file__).parent / "NotoSansSC.ttf"),
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def render_pdf(content: str, output_path: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()

    font_path = _find_chinese_font()
    if font_path:
        pdf.add_font("CJK", "", font_path)
        pdf.add_font("CJK", "B", font_path)
        title_font = ("CJK", "B", 16)
        body_font = ("CJK", "", 10)
        heading_font = ("CJK", "B", 13)
    else:
        print("  [警告] 未找到中文字体，PDF 中文将无法显示")
        title_font = ("Helvetica", "B", 16)
        body_font = ("Helvetica", "", 10)
        heading_font = ("Helvetica", "B", 13)

    body_w = pdf.w - pdf.l_margin - pdf.r_margin

    for line in content.split("\n"):
        line = line.rstrip()
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue

        pdf.set_x(pdf.l_margin)  # reset x before every line

        if stripped.startswith("# ") or stripped.startswith("《"):
            pdf.set_font(*title_font)
            pdf.multi_cell(body_w, 10, stripped.lstrip("# "), align="C")
            pdf.ln(4)
        elif stripped.startswith("## "):
            pdf.set_font(*heading_font)
            pdf.ln(3)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(body_w, 7, stripped.lstrip("# ").lstrip("## "))
            pdf.ln(1)
        elif stripped.startswith("### "):
            pdf.set_font(*heading_font)
            pdf.multi_cell(body_w, 7, stripped.lstrip("# ").lstrip("### "))
            pdf.ln(1)
        elif stripped.startswith("|") and "|" in stripped[1:]:
            pdf.set_font(*body_font)
            pdf.multi_cell(body_w, 5, stripped)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font(*body_font)
            pdf.set_x(pdf.l_margin + 4)
            pdf.multi_cell(body_w - 4, 5, f"· {stripped[2:]}")
        elif stripped[0].isdigit() and ("." in stripped[:4] or "、" in stripped[:4]):
            pdf.set_font(*body_font)
            pdf.multi_cell(body_w, 5, stripped)
        else:
            pdf.set_font(*body_font)
            pdf.multi_cell(body_w, 5, stripped)

    pdf.output(str(output_path))
    print(f"  PDF: {output_path.name} ({output_path.stat().st_size} bytes)")


def render_docx(content: str, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)

    for line in content.split("\n"):
        line = line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# ") or stripped.startswith("《"):
            p = doc.add_heading(stripped.lstrip("# "), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped.lstrip("# ").lstrip("## "), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped.lstrip("# ").lstrip("### "), level=2)
        elif stripped.startswith("|") and "|" in stripped[1:]:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if any(c.startswith("---") or c.startswith(":--") for c in cells):
                continue
            if not hasattr(doc, '_last_table'):
                doc._last_table = doc.add_table(rows=1, cols=len(cells))
                doc._last_table_cols = len(cells)
                doc._last_table.style = 'Light Grid Accent 1'
                for i, cell_text in enumerate(cells):
                    doc._last_table.rows[0].cells[i].text = cell_text
                    for p in doc._last_table.rows[0].cells[i].paragraphs:
                        for run in p.runs:
                            run.bold = True
            else:
                row = doc._last_table.add_row()
                for i in range(min(len(cells), doc._last_table_cols)):
                    row.cells[i].text = cells[i]
        else:
            if hasattr(doc, '_last_table'):
                del doc._last_table
                del doc._last_table_cols
            if stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                doc.add_paragraph(stripped)

    doc.save(str(output_path))
    print(f"  DOCX: {output_path.name} ({output_path.stat().st_size} bytes)")


def render_xlsx(content: str, output_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    sheets = re.split(r'^\[Sheet:\s*(.+?)\]$', content, flags=re.MULTILINE)
    # sheets[0] = text before first sheet, sheets[1] = name, sheets[2] = content, ...
    if len(sheets) <= 1:
        print(f"  [警告] XLSX 无 Sheet 标记，跳过: {output_path.name}")
        return

    # Skip prelude text
    start = 1 if not sheets[0].strip() else 2
    if start == 2:
        sheet_pairs = list(zip(sheets[1::2], sheets[2::2]))
    else:
        # First token is preamble, rest are name/content pairs
        pairs_raw = sheets[1:]
        sheet_pairs = list(zip(pairs_raw[0::2], pairs_raw[1::2]))

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    header_font = Font(bold=True)

    for sheet_name, sheet_content in sheet_pairs:
        ws = wb.create_sheet(title=sheet_name[:31])  # Excel max sheet name
        for row_idx, line in enumerate(sheet_content.strip().split("\n"), 1):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            cells = [c.strip().strip('"') for c in line.split(",")]
            for col_idx, val in enumerate(cells, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                if row_idx == 1:
                    cell.font = header_font
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True)

        # Auto-width
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    wb.save(str(output_path))
    print(f"  XLSX: {output_path.name} ({output_path.stat().st_size} bytes, {len(sheet_pairs)} sheets)")


def render_html(content: str, output_path: Path) -> None:
    html_body = content
    # Very simple markdown→html conversion for the demo
    for i in range(4, 0, -1):
        prefix = "#" * i + " "
        html_body = html_body.replace(prefix, f"<h{i}>") + f"</h{i}>"
    html_body = html_body.replace("\n\n", "</p>\n<p>")

    template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>REST API 接口规范 — IWMS 项目</title>
<style>
  body {{ font-family: -apple-system, 'Microsoft YaHei', sans-serif; max-width: 960px; margin: 0 auto; padding: 20px; line-height: 1.7; color: #333; }}
  h1 {{ border-bottom: 2px solid #1a73e8; padding-bottom: 8px; }}
  h2 {{ border-bottom: 1px solid #e0e0e0; padding-bottom: 6px; margin-top: 32px; }}
  h3 {{ margin-top: 24px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  th {{ background: #f5f5f5; font-weight: 600; }}
  code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 3px; font-size: 13px; }}
  pre {{ background: #f8f8f8; padding: 16px; border-radius: 6px; overflow-x: auto; }}
</style>
</head>
<body>
<p>{html_body}</p>
</body>
</html>"""
    output_path.write_text(template, encoding="utf-8")
    print(f"  HTML: {output_path.name} ({output_path.stat().st_size} bytes)")


def render_md(content: str, output_path: Path) -> None:
    output_path.write_text(content, encoding="utf-8")
    print(f"  MD: {output_path.name} ({output_path.stat().st_size} bytes)")


def render_txt(content: str, output_path: Path, encoding: str = "utf-8") -> None:
    output_path.write_text(content, encoding=encoding)
    print(f"  TXT: {output_path.name} ({output_path.stat().st_size} bytes, {encoding})")


FORMAT_HANDLERS = {
    "pdf":  lambda c, p, **kw: render_pdf(c, p),
    "docx": lambda c, p, **kw: render_docx(c, p),
    "xlsx": lambda c, p, **kw: render_xlsx(c, p),
    "html": lambda c, p, **kw: render_html(c, p),
    "md":   lambda c, p, **kw: render_md(c, p),
    "txt":  lambda c, p, **kw: render_txt(c, p, encoding=kw.get("encoding", "utf-8")),
}


# ---- 主流程 ----

def parse_source_file(source_path: Path) -> dict:
    """解析 markdown 源文件，提取 frontmatter 和正文内容。"""
    text = source_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    meta = {}
    content_start = 0

    # 解析顶部的 key: value 行
    for i, line in enumerate(lines):
        if ":" in line and not line.strip().startswith("#") and not line.strip().startswith("-"):
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip()
            if key in ("format", "output", "encoding"):
                meta[key] = value
            content_start = i + 1
        else:
            break

    # 正文从第一个非 meta 行开始
    meta["content"] = "\n".join(lines[content_start:]).strip()
    return meta


def main():
    source_dir = Path(__file__).parent / "test_docs" / "iwms" / "source"
    output_dir = Path(__file__).parent / "test_docs" / "iwms" / "output"

    if not source_dir.exists():
        print(f"源文件目录不存在: {source_dir}")
        print("请先创建源文件，格式: format: <fmt>\\noutput: <filename>\\n\\n<content>")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    source_files = sorted(source_dir.glob("*.md"))
    if not source_files:
        print(f"源文件目录为空: {source_dir}")
        return

    print(f"找到 {len(source_files)} 个源文件\n")

    for src_path in source_files:
        print(f"渲染: {src_path.name} ...")
        meta = parse_source_file(src_path)
        fmt = meta.get("format", "md")
        output_name = meta.get("output", src_path.stem + "." + fmt)
        content = meta["content"]
        encoding = meta.get("encoding", "utf-8")

        if not content:
            print(f"  [跳过] 内容为空")
            continue

        handler = FORMAT_HANDLERS.get(fmt)
        if not handler:
            print(f"  [跳过] 不支持的格式: {fmt}")
            continue

        out_path = output_dir / output_name
        handler(content, out_path, encoding=encoding)

    print(f"\n全部完成！输出目录: {output_dir.resolve()}")
    print("文件列表:")
    for f in sorted(output_dir.iterdir()):
        if f.is_file():
            size_kb = f.stat().st_size / 1024
            print(f"  {f.name} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
