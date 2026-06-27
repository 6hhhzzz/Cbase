"""DOCX 解析器 — 使用 python-docx 提取文本和表格。"""

from docx import Document

from common import get_logger
from models.document import ParseResult
from etl.common.table_utils import rows_to_markdown
from .base import BaseParser

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """DOCX 文件解析器。提取段落文本和表格数据。"""

    async def _do_parse(self, file_path: str) -> ParseResult:
        doc = Document(file_path)
        parts: list[str] = []
        tables: list[dict] = []
        table_count = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = _find_paragraph(doc, element)
                if para and para.text.strip():
                    parts.append(para.text.strip())
            elif tag == "tbl":
                table = _find_table(doc, element)
                if table:
                    table_data = _extract_table(table)
                    if table_data:
                        table_count += 1
                        tables.append({
                            "index": table_count,
                            "headers": table_data[0] if table_data else [],
                            "rows": table_data[1:] if len(table_data) > 1 else [],
                        })
                        parts.append(rows_to_markdown(table_data))

        raw_text = "\n\n".join(parts)
        logger.info(f"DOCX 解析完成: {file_path}, chars={len(raw_text)}, tables={table_count}")
        return ParseResult(
            file_type="docx", raw_text=raw_text, page_count=1,
            tables=tables, metadata={"table_count": table_count},
        )

    def supports(self, file_type: str) -> bool:
        return file_type == "docx"


def _find_paragraph(doc, element):
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _extract_table(table) -> list[list[str]]:
    data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(c for c in cells):
            data.append(cells)
    return data
