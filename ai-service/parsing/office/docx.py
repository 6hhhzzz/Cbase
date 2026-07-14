"""DOCX 解析器 — 增强版：保留表格结构 + 提取表格内段落。

相比旧版改进：
    - 表格不再序列化为 Markdown 字符串混入文本流
    - 表格独立为 TableBlock，含 HTML 和语义描述
    - 保留段落样式信息用于标题检测
    - v2: 递归提取表格单元格内的段落（修复内容丢失 bug）
"""

from docx import Document

from common import get_logger
from common.utils import estimate_tokens
from parsing.base import BaseParser
from parsing.models import ParsedDocument, DocumentMetadata, TextBlock, TableBlock

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """DOCX 文件解析器。提取段落文本和表格数据为结构化 blocks。"""

    async def _do_parse(self, file_path: str) -> ParsedDocument:
        doc = Document(file_path)
        blocks: list[TextBlock | TableBlock] = []
        table_count = 0

        # 1. 顶层段落
        for para in doc.paragraphs:
            if para.text.strip():
                level = _detect_heading_level(para)
                layout_type = "title" if level else "text"
                blocks.append(TextBlock(
                    text=para.text.strip(),
                    page_num=0,
                    layout_type=layout_type,
                    level=level,
                ))

        # 2. 表格单元格内的段落（关键修复：之前遗漏了嵌套在表格中的内容）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            level = _detect_heading_level(para)
                            layout_type = "title" if level else "text"
                            blocks.append(TextBlock(
                                text=para.text.strip(),
                                page_num=0,
                                layout_type=layout_type,
                                level=level,
                            ))

        # 3. 表格结构信息（供 ContextEnricher 注入上下文）
        for table in doc.tables:
            table_data = _extract_table(table)
            if table_data:
                table_count += 1
                html = _to_html_table(table_data)
                description = _describe_table(table_data, table_count)
                blocks.append(TableBlock(
                    html=html,
                    description=description,
                    page_num=0,
                    caption=None,
                    metadata={"table_index": table_count},
                ))

        logger.info(
            f"DOCX 解析完成: {file_path}, text_blocks={len([b for b in blocks if isinstance(b, TextBlock)])}, "
            f"tables={table_count}"
        )

        return ParsedDocument(
            blocks=blocks,
            metadata=DocumentMetadata(
                file_type="docx",
                file_name=file_path,
                page_count=1,
                title=_extract_doc_title(doc),
            ),
        )

    def supports(self, file_type: str) -> bool:
        return file_type == "docx"


# ---- 内部辅助 ----

def _detect_heading_level(para) -> int | None:
    """检测段落是否为标题及其级别。"""
    style_name = (para.style.name if para.style else "").lower()
    if "heading" in style_name or "标题" in style_name:
        try:
            for part in style_name.split():
                if part.isdigit():
                    return min(int(part), 6)
        except (ValueError, TypeError):
            pass
        return 1
    return None


def _extract_doc_title(doc) -> str | None:
    """尝试提取文档标题。"""
    if doc.core_properties.title:
        return doc.core_properties.title
    for para in doc.paragraphs:
        if para.style and "heading" in (para.style.name or "").lower():
            return para.text.strip()[:200]
    return None


def _extract_table(table) -> list[list[str]]:
    """从 DOCX Table 提取二维数组。"""
    data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(c for c in cells):
            data.append(cells)
    return data


def _to_html_table(data: list[list[str]]) -> str:
    """二维数组 → HTML <table>。"""
    if not data:
        return ""
    rows = ["<table>"]
    rows.append("<thead><tr>")
    for cell in data[0]:
        rows.append(f"<th>{cell}</th>")
    rows.append("</tr></thead>")
    if len(data) > 1:
        rows.append("<tbody>")
        for row in data[1:]:
            rows.append("<tr>")
            for cell in row:
                rows.append(f"<td>{cell}</td>")
            rows.append("</tr>")
        rows.append("</tbody>")
    rows.append("</table>")
    return "\n".join(rows)


def _describe_table(data: list[list[str]], index: int) -> str:
    """生成表格语义描述。"""
    if not data:
        return ""
    headers = data[0]
    row_count = len(data) - 1
    col_count = len(headers)
    token_est = estimate_tokens("\n".join(" ".join(r) for r in data))
    return (
        f"表格 {index}：{row_count} 行 × {col_count} 列"
        f"，表头为 [{', '.join(headers[:5])}]"
        f"，约 {token_est} tokens"
    )
